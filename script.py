
from docx import Document
import pandas as pd
import re

df = pd.read_excel('planilha_dados.xlsx')

doc = Document('modelo_documento.docx')

def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(str(replace), inline[i].text)
                    inline[i].text = text

for index, row in df.iterrows():
    filename = "modelo_documento.docx"
    doc = Document(filename)
    for column_name in df.columns:
        regex = re.compile("{" + column_name + "}")
        replace = row[column_name]
        docx_replace_regex(doc, regex , replace)
    doc.save('documentos_gerados/' + str(index) + '.docx')
